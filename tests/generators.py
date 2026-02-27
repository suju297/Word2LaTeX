"""Generate DOCX samples for testing."""
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SAMPLES_DIR = Path(__file__).parent / "samples"

def ensure_samples_dir():
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)

def generate_text_formatting():
    """Generate sample with basic text formatting."""
    doc = docx.Document()
    doc.add_heading('Text Formatting Test', 0)

    p = doc.add_paragraph('This is a ')
    p.add_run('bold').bold = True
    p.add_run(' word.')

    p = doc.add_paragraph('This is an ')
    p.add_run('italic').italic = True
    p.add_run(' word.')

    p = doc.add_paragraph('This is ')
    p.add_run('underlined').underline = True
    p.add_run('.')

    p = doc.add_paragraph('Mixed: ')
    r = p.add_run('Bold and Italic')
    r.bold = True
    r.italic = True
    
    doc.save(SAMPLES_DIR / "feature_text.docx")

def generate_lists():
    """Generate sample with lists."""
    doc = docx.Document()
    doc.add_heading('List Test', 0)

    doc.add_paragraph('Unordered List:')
    doc.add_paragraph('Item 1', style='List Bullet')
    doc.add_paragraph('Item 2', style='List Bullet')
    
    doc.add_paragraph('Ordered List:')
    doc.add_paragraph('First', style='List Number')
    doc.add_paragraph('Second', style='List Number')

    doc.save(SAMPLES_DIR / "feature_lists.docx")

def generate_tables():
    """Generate sample with tables."""
    doc = docx.Document()
    doc.add_heading('Table Test', 0)

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"

    doc.save(SAMPLES_DIR / "feature_tables.docx")

def generate_all():
    ensure_samples_dir()
    generate_text_formatting()
    generate_lists()
    generate_tables()
    print(f"Generated samples in {SAMPLES_DIR}")

if __name__ == "__main__":
    generate_all()
